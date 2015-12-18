####################################################
#                  Revision: 1.0                   #
#              Updated on: 11/06/2015              #
####################################################

####################################################
#                                                  #
#   This file writes/append data to the Template   #
#   Word File.                                     #
#                                                  #
#   Author: Zankar Sanghavi                        #
#                                                  #
#   © Dot Hill Systems Corporation                 #
#                                                  #
####################################################

import os
import sys
###################################
#  Importing from other Directory
###################################
import sys
os.chdir('..')
c_path = os.getcwd()
sys.path.insert(0, r''+str(c_path)+'/Common Scripts')
import user_inputs_ICS
import fixed_data_ICS
import report_functions
import extract_lists
import modify_word_docx

sys.path.insert(0, r''+str(c_path)+'/IO Stress')
import log_functions


###################################
#  Importing from Current Directory
###################################
sys.path.insert(0, r''+str(c_path)+'/SDR')
import pandas
import csv
import numpy as np

import extract_f2_f3_sdr
import time
import zipfile

from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_TABLE_DIRECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.shared import Cm
import re


###################################
#  To call functions from different
#  files. 
###################################
lf=log_functions.Log_Functions
ed= extract_f2_f3_sdr.Extract_F2_F3_SDR
ui= user_inputs_ICS.User_Inputs
fd= fixed_data_ICS.Fixed_Data
rf= report_functions.Report_Functions

class Generate_Word_SDR:

    ####################################################
    #   This Function does:
    #   
    #   1.) Uses "unzip_pull_log" function to unzip and
    #   find the .logs file in the folder to read the 
    #   data.
    #
    #   2.) Checks Errors, extracts: Hardware, Software
    #   and PhyStats/DPSLD Informations. Generates two
    #   DPSLD files to write in WORD report. 
    #   
    #   3.) It writes required data, with formatting
    #   into a WORD report. 
    #   
    ####################################################
    def generate_final_report(file_name, file_name1, fixed_dir,
        part_no, test_name, no_of_files, csv_file_list, log_list,
        fw_no_names, chassis_names, cntrllr_names, t_fw_type):
        
        document = Document(r''+str(fixed_dir)+'\\'+str(part_no)
                            +str(test_name)+'.docx')
        
        document.add_page_break()
        f=no_of_files
        error_collection=[]
        for f in range(no_of_files):
            progress=(round((float(100/int(no_of_files))*f),2)) 
            # to show Report Progress
            
            print('\nReport Progess: ',progress,'%\n')
            
            log_temp = str(no_of_files)+' - Log file'
            log_data= lf.unzip_pull_log(log_list[f], 'store', log_temp) 
            #pull .logs file from zip folder
            
            
            ###################################
            #  Error check and data extraction
            ###################################
            [write_sum, read_sum, iter_flag, hw_list, host_list
                , sasmap_list]= ed.generate_f2_f3_sdr(
                                csv_file_list[f], log_data)

            
            ###################
            #   Summary
            ###################
            #document.add_page_break()
            
            document.add_heading('Shutdown Reboot test Summary'\
                                 'for ' + str(fw_no_names[f]) + 
                                '\\' + str(chassis_names[f]) +  
                                '\\'  + str(cntrllr_names[f]) + 
                                ' chassis' ,level=3)
            
            temp_para=document.add_paragraph()
            
            paragraph_format = temp_para.paragraph_format
            paragraph_format.left_indent
            paragraph_format.left_indent = Inches(0.5)
            
            run = temp_para.add_run('Read error(s): '
                                    +str(read_sum)+ 
                                    '\nWrite error(s): '
                                    +str(write_sum)
                                    + '\nNumber of Iterations: '
                                    +str(iter_flag)+'\n')
            font = run.font
            font.name = 'Courier New'
            font.size = Pt(11)
            
            
            ###################
            #   F2 Menu
            ###################
            document.add_heading('Shutdown Reboot test (F2 menu)'\
                                 'in ' + str(fw_no_names[f]) + 
                                 '\\' + str(chassis_names[f]) +
                                 '\\'  + str(cntrllr_names[f]) + 
                                 ' chassis' ,level=3)
            
            for i in range(len(hw_list)):
                temp_para= document.add_paragraph()
                temp_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = temp_para.add_run(hw_list[i])
                font = run.font
                font.name = 'Courier New'
                font.size = Pt(11)
                             
            
            ###################
            #   F3 Menu
            ###################
            document.add_page_break()
            
            document.add_heading('Shutdown Reboot test (F3 menu)'\
                                 'in ' + str(fw_no_names[f]) + 
                                 '\\' + str(chassis_names[f]) +
                                 '\\'  + str(cntrllr_names[f]) +
                                 ' chassis',level=3)
            
            
            for i in range(len(sasmap_list)):
                temp_para1= document.add_paragraph() 
                run = temp_para1.add_run(sasmap_list[i])
                font = run.font
                font.name = 'Courier New'
                font.size = Pt(11)

            document.save(r''+str(fixed_dir)+'\\'+str(part_no)
                          +str(test_name)+'.docx')
        return error_collection
        
        
#####################################
#              END                  #
#####################################